from pathlib import Path
from psycopg import Connection
from docx import Document

from app.core.config import get_settings
from app.core.exceptions import AppError
from app.services.common import page_bounds, pagination


def create_asset(conn: Connection, payload: dict) -> dict:
    parent_asset_id = payload.get("parent_asset_id")
    root_asset_id = payload.get("root_asset_id")
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO research_assets (
                asset_type, title, content_markdown, summary, company_code, industry, task_id, parent_asset_id, root_asset_id
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING *
            """,
            (
                payload["asset_type"], payload["title"], payload["content_markdown"],
                payload.get("summary"), payload.get("company_code"), payload.get("industry"),
                payload.get("task_id"), parent_asset_id, root_asset_id,
            ),
        )
        row = cur.fetchone()
        if parent_asset_id and not root_asset_id:
            cur.execute(
                """
                UPDATE research_assets
                SET root_asset_id = COALESCE((SELECT root_asset_id FROM research_assets WHERE id = %s), %s)
                WHERE id = %s
                RETURNING *
                """,
                (parent_asset_id, parent_asset_id, row["id"]),
            )
            row = cur.fetchone()
        cur.execute(
            "INSERT INTO asset_revisions (asset_id, version, content_markdown, summary, change_note) VALUES (%s, 1, %s, %s, 'initial')",
            (row["id"], row["content_markdown"], row.get("summary")),
        )
        _replace_asset_sources(cur, str(row["id"]), payload.get("sources", []))
        _replace_asset_tags(cur, str(row["id"]), payload.get("tag_ids", []))
    return get_asset(conn, str(row["id"]))


def list_assets(
    conn: Connection,
    page: int,
    page_size: int,
    asset_type: str | None = None,
    company_code: str | None = None,
    status: str | None = None,
    tag_id: str | None = None,
    tag_code: str | None = None,
    keyword: str | None = None,
) -> dict:
    limit, offset = page_bounds(page, page_size)
    where = ["ra.deleted_at IS NULL"]
    params: list = []
    for field, value in (("asset_type", asset_type), ("company_code", company_code), ("status", status)):
        if value:
            where.append(f"ra.{field} = %s")
            params.append(value)
    if keyword:
        where.append("(ra.title ILIKE %s OR ra.summary ILIKE %s)")
        params.extend([f"%{keyword}%", f"%{keyword}%"])
    if tag_id:
        where.append("EXISTS (SELECT 1 FROM research_asset_tags rat WHERE rat.asset_id = ra.id AND rat.tag_id = %s)")
        params.append(tag_id)
    if tag_code:
        where.append(
            """
            EXISTS (
                SELECT 1
                FROM research_asset_tags rat
                JOIN tags t ON t.id = rat.tag_id
                WHERE rat.asset_id = ra.id AND t.tag_code = %s
            )
            """
        )
        params.append(tag_code)
    with conn.cursor() as cur:
        cur.execute(f"SELECT COUNT(*) AS count FROM research_assets ra WHERE {' AND '.join(where)}", tuple(params))
        total = cur.fetchone()["count"]
        cur.execute(
            f"""
            SELECT
                ra.*,
                COUNT(DISTINCT src.id) AS source_count,
                COALESCE(array_remove(array_agg(DISTINCT t.tag_name), NULL), ARRAY[]::varchar[]) AS tag_names
            FROM research_assets ra
            LEFT JOIN asset_sources src ON src.asset_id = ra.id
            LEFT JOIN research_asset_tags rat ON rat.asset_id = ra.id
            LEFT JOIN tags t ON t.id = rat.tag_id
            WHERE {' AND '.join(where)}
            GROUP BY ra.id
            ORDER BY ra.created_at DESC
            LIMIT %s OFFSET %s
            """,
            tuple(params + [limit, offset]),
        )
        rows = cur.fetchall()
    return {"items": [dict(row) for row in rows], "pagination": pagination(page, limit, total)}


def get_asset(conn: Connection, asset_id: str) -> dict:
    with conn.cursor() as cur:
        cur.execute("SELECT * FROM research_assets WHERE id = %s AND deleted_at IS NULL", (asset_id,))
        row = cur.fetchone()
    if not row:
        raise AppError(5001, "asset not found", 404)
    asset = dict(row)
    asset["sources"] = list_asset_sources(conn, asset_id)["items"]
    asset["tags"] = list_asset_tags(conn, asset_id)["items"]
    return asset


def update_asset(conn: Connection, asset_id: str, payload: dict) -> dict:
    current = get_asset(conn, asset_id)
    new_version = current["version"] + 1
    with conn.cursor() as cur:
        cur.execute(
            """
            UPDATE research_assets
            SET title = COALESCE(%s, title),
                content_markdown = COALESCE(%s, content_markdown),
                summary = COALESCE(%s, summary),
                status = COALESCE(%s, status),
                version = %s,
                updated_at = NOW()
            WHERE id = %s
            RETURNING *
            """,
            (
                payload.get("title"), payload.get("content_markdown"), payload.get("summary"),
                payload.get("status"), new_version, asset_id,
            ),
        )
        row = cur.fetchone()
        cur.execute(
            "INSERT INTO asset_revisions (asset_id, version, content_markdown, summary, change_note) VALUES (%s, %s, %s, %s, %s)",
            (asset_id, new_version, row["content_markdown"], row.get("summary"), payload.get("change_note")),
        )
        if "sources" in payload:
            _replace_asset_sources(cur, asset_id, payload.get("sources", []))
        if "tag_ids" in payload:
            _replace_asset_tags(cur, asset_id, payload.get("tag_ids", []))
    return get_asset(conn, asset_id)


def list_revisions(conn: Connection, asset_id: str) -> dict:
    get_asset(conn, asset_id)
    with conn.cursor() as cur:
        cur.execute("SELECT * FROM asset_revisions WHERE asset_id = %s ORDER BY version DESC", (asset_id,))
        rows = cur.fetchall()
    return {"items": [dict(row) for row in rows]}


def list_asset_sources(conn: Connection, asset_id: str) -> dict:
    with conn.cursor() as cur:
        cur.execute("SELECT * FROM asset_sources WHERE asset_id = %s ORDER BY created_at ASC", (asset_id,))
        rows = cur.fetchall()
    return {"items": [dict(row) for row in rows]}


def list_asset_tags(conn: Connection, asset_id: str) -> dict:
    with conn.cursor() as cur:
        cur.execute(
            """
            SELECT t.*
            FROM tags t
            JOIN research_asset_tags rat ON rat.tag_id = t.id
            WHERE rat.asset_id = %s
            ORDER BY t.tag_type, t.tag_name
            """,
            (asset_id,),
        )
        rows = cur.fetchall()
    return {"items": [dict(row) for row in rows]}


def export_asset(conn: Connection, asset_id: str, export_format: str = "markdown") -> dict:
    asset = get_asset(conn, asset_id)
    export_dir = Path(get_settings().export_dir)
    export_dir.mkdir(parents=True, exist_ok=True)
    export_format = export_format.lower()
    if export_format in ("markdown", "md"):
        file_name = _build_export_file_name(asset, "md")
        path = export_dir / file_name
        path.write_text(asset["content_markdown"], encoding="utf-8")
        return _build_export_response(asset_id, "markdown", path)
    if export_format == "docx":
        file_name = _build_export_file_name(asset, "docx")
        path = export_dir / file_name
        _write_docx(path, asset["title"], asset["content_markdown"])
        return _build_export_response(asset_id, "docx", path)
    raise AppError(5003, "asset export format not supported", 400)


def _write_docx(path: Path, title: str, markdown: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:3].endswith(". ") and line[0].isdigit():
            doc.add_paragraph(line[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(path)


def _replace_asset_sources(cur, asset_id: str, sources: list[dict]) -> None:
    cur.execute("DELETE FROM asset_sources WHERE asset_id = %s", (asset_id,))
    for item in sources:
        cur.execute(
            """
            INSERT INTO asset_sources (asset_id, source_type, source_ref_id, source_id_text, note)
            VALUES (%s, %s, %s, %s, %s)
            """,
            (
                asset_id,
                item["source_type"],
                item["source_ref_id"],
                item.get("source_id_text"),
                item.get("note"),
            ),
        )


def _replace_asset_tags(cur, asset_id: str, tag_ids: list[str]) -> None:
    cur.execute("DELETE FROM research_asset_tags WHERE asset_id = %s", (asset_id,))
    for tag_id in tag_ids:
        cur.execute(
            "INSERT INTO research_asset_tags (asset_id, tag_id) VALUES (%s, %s) ON CONFLICT DO NOTHING",
            (asset_id, tag_id),
        )


def _build_export_file_name(asset: dict, extension: str) -> str:
    asset_type = (asset.get("asset_type") or "asset").strip().lower()
    company_code = (asset.get("company_code") or "general").strip().lower()
    created_at = asset.get("created_at")
    date_part = created_at.strftime("%Y%m%d") if created_at else "unknown"
    return f"{asset_type}_{company_code}_{date_part}.{extension}"


def _build_export_response(asset_id: str, export_format: str, path: Path) -> dict:
    return {
        "asset_id": asset_id,
        "format": export_format,
        "file_name": path.name,
        "download_path": str(path),
        "file_path": str(path),
    }
