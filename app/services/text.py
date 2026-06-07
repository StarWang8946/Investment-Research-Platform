import re
from pathlib import Path


HEADING_RE = re.compile(
    r"^\s*(#{1,6}\s+\S+|第[一二三四五六七八九十百\d]+[章节部分][：:、.\s]?.+|\d+(?:\.\d+){0,4}[、.\s]+.+)$"
)
PAGE_MARK_RE = re.compile(r"^\[page \d+\]$", re.IGNORECASE)
SENTENCE_SPLIT_RE = re.compile(r"(?<=[。！？!?；;])")


def read_text_file(path: str) -> str:
    data = open(path, "rb").read()
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def read_document_text(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return read_pdf_text(path)
    if suffix == ".docx":
        return read_docx_text(path)
    return read_text_file(path)


def read_pdf_text(path: str) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("pymupdf is required to parse PDF files") from exc

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page_no, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                parts.append(f"[page {page_no}]\n{text}")
    return "\n\n".join(parts)


def read_docx_text(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def split_chunks(text: str, chunk_size: int = 800, chunk_overlap: int = 120) -> list[str]:
    normalized = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not normalized:
        return []
    if chunk_size <= 0:
        chunk_size = 800
    chunk_overlap = max(0, min(chunk_overlap, chunk_size // 2))

    blocks = _semantic_blocks(normalized, chunk_size)
    return _pack_blocks(blocks, chunk_size, chunk_overlap)


def _semantic_blocks(text: str, chunk_size: int) -> list[str]:
    blocks: list[str] = []
    current_heading: str | None = None
    current_page: str | None = None

    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = _normalize_paragraph(paragraph)
        if not paragraph:
            continue

        if PAGE_MARK_RE.match(paragraph):
            current_page = paragraph
            continue

        if _is_heading(paragraph):
            current_heading = paragraph
            continue

        for part in _split_oversized_paragraph(paragraph, chunk_size):
            blocks.append(_with_context(part, current_page, current_heading))

    return blocks


def _normalize_paragraph(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ""
    if len(lines) == 1:
        return lines[0]
    if any(PAGE_MARK_RE.match(line) or _is_heading(line) for line in lines):
        return "\n".join(lines)
    return " ".join(lines)


def _is_heading(text: str) -> bool:
    if "\n" in text or len(text) > 80:
        return False
    return bool(HEADING_RE.match(text))


def _split_oversized_paragraph(paragraph: str, chunk_size: int) -> list[str]:
    if len(paragraph) <= chunk_size:
        return [paragraph]

    sentences = [part.strip() for part in SENTENCE_SPLIT_RE.split(paragraph) if part.strip()]
    if len(sentences) <= 1:
        return _split_by_length(paragraph, chunk_size)

    parts: list[str] = []
    current = ""
    for sentence in sentences:
        if len(sentence) > chunk_size:
            if current:
                parts.append(current)
                current = ""
            parts.extend(_split_by_length(sentence, chunk_size))
            continue
        candidate = sentence if not current else current + sentence
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            parts.append(current)
            current = sentence
    if current:
        parts.append(current)
    return parts


def _split_by_length(text: str, chunk_size: int) -> list[str]:
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end
    return chunks


def _with_context(text: str, page: str | None, heading: str | None) -> str:
    prefixes = []
    if page and page not in text:
        prefixes.append(page)
    if heading and heading not in text:
        prefixes.append(heading)
    if not prefixes:
        return text
    return "\n".join(prefixes + [text])


def _pack_blocks(blocks: list[str], chunk_size: int, chunk_overlap: int) -> list[str]:
    chunks: list[str] = []
    current_parts: list[str] = []
    current_len = 0

    for block in blocks:
        separator_len = 2 if current_parts else 0
        if current_parts and current_len + separator_len + len(block) > chunk_size:
            chunks.append("\n\n".join(current_parts))
            current_parts = _overlap_tail(current_parts, chunk_overlap)
            current_len = len("\n\n".join(current_parts)) if current_parts else 0
            separator_len = 2 if current_parts else 0
        if current_parts and current_len + separator_len + len(block) > chunk_size:
            chunks.append("\n\n".join(current_parts))
            current_parts = []
            current_len = 0
        current_parts.append(block)
        current_len = len("\n\n".join(current_parts))

    if current_parts:
        chunks.append("\n\n".join(current_parts))
    return chunks


def _overlap_tail(parts: list[str], chunk_overlap: int) -> list[str]:
    if chunk_overlap <= 0:
        return []
    selected: list[str] = []
    total = 0
    for part in reversed(parts):
        part_len = len(part) + (2 if selected else 0)
        if selected and total + part_len > chunk_overlap:
            break
        selected.append(part)
        total += part_len
        if total >= chunk_overlap:
            break
    return list(reversed(selected))


def estimate_tokens(text: str) -> int:
    return max(1, len(text) // 2)
