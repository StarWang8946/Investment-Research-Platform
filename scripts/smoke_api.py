from pathlib import Path
import sys
from uuid import uuid4

from fastapi.testclient import TestClient
import psycopg
from psycopg.rows import dict_row

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.core.config import get_settings
from app.main import app


SAMPLE_TEXT = (
    "贵州茅台收入增长主要来自直营渠道扩张和产品结构优化。\n"
    "公司现金流稳定，毛利率保持高位。\n"
    "白酒行业需求韧性较强，但渠道库存需要持续观察。"
)


def unwrap(response):
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["code"] == 0, payload
    assert payload["message"] == "ok", payload
    assert "data" in payload, payload
    return payload["data"]


def make_samples(base_dir: Path, run_id: str) -> list[Path]:
    base_dir.mkdir(parents=True, exist_ok=True)
    sample_text = f"{SAMPLE_TEXT}\n本次测试批次：{run_id}"
    txt_path = base_dir / f"sample_report_{run_id}.txt"
    txt_path.write_text(sample_text, encoding="utf-8")

    docx_path = base_dir / f"sample_report_{run_id}.docx"
    from docx import Document

    doc = Document()
    doc.add_heading("贵州茅台测试研报", level=1)
    doc.add_paragraph(sample_text)
    doc.save(docx_path)

    pdf_path = base_dir / f"sample_report_{run_id}.pdf"
    import fitz

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), sample_text)
    pdf.save(pdf_path)
    pdf.close()

    return [txt_path, docx_path, pdf_path]


def upload_and_ingest(client: TestClient, path: Path) -> str:
    with path.open("rb") as file:
        uploaded = unwrap(
            client.post(
                "/api/v1/documents",
                data={
                    "doc_type": "report",
                    "title": f"Smoke {path.suffix} {uuid4().hex[:8]}",
                    "company_code": "600519",
                    "company_name": "贵州茅台",
                },
                files={"file": (path.name, file, "application/octet-stream")},
            )
        )
    document_id = uploaded["id"]
    ingested = unwrap(
        client.post(
            f"/api/v1/documents/{document_id}/ingest",
            json={"chunk_size": 80, "chunk_overlap": 10, "force_reingest": True},
        )
    )
    assert ingested["status"] == "parsed", ingested
    assert ingested["chunk_count"] >= 1, ingested

    chunks = unwrap(client.get(f"/api/v1/documents/{document_id}/chunks"))
    assert chunks["pagination"]["total"] >= 1, chunks
    return document_id


def assert_citations_saved(task_id: str, expected_count: int) -> None:
    with psycopg.connect(get_settings().database_url, row_factory=dict_row) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT COUNT(*) AS count, ARRAY_AGG(citation_no ORDER BY citation_no) AS citation_nos
                FROM citations
                WHERE task_id = %s
                """,
                (task_id,),
            )
            row = cur.fetchone()
    assert row["count"] == expected_count, row
    assert row["citation_nos"] == list(range(1, expected_count + 1)), row


def assert_asset_citations_saved(asset_id: str, expected_count: int) -> None:
    with psycopg.connect(get_settings().database_url, row_factory=dict_row) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT COUNT(*) AS count, ARRAY_AGG(citation_no ORDER BY citation_no) AS citation_nos
                FROM citations
                WHERE asset_id = %s
                """,
                (asset_id,),
            )
            row = cur.fetchone()
    assert row["count"] == expected_count, row
    assert row["citation_nos"] == list(range(1, expected_count + 1)), row


def main() -> None:
    client = TestClient(app)
    run_id = uuid4().hex[:8]

    health = client.get("/health")
    assert health.status_code == 200, health.text
    assert health.json()["status"] == "ok", health.text

    system = unwrap(client.get("/api/v1/system/info"))
    assert system["vector_enabled"] is True, system

    code = f"SMOKE{uuid4().hex[:8]}"
    company = unwrap(
        client.post(
            "/api/v1/companies",
            json={
                "company_code": code,
                "company_name": f"Smoke Test Company {code}",
                "company_short_name": code,
                "market": "TEST",
                "is_active": True,
            },
        )
    )
    assert company["company_code"] == code, company
    companies = unwrap(client.get("/api/v1/companies", params={"market": "TEST", "keyword": code}))
    assert any(item["company_code"] == code for item in companies["items"]), companies

    sample_paths = make_samples(Path("/private/tmp/investment_research_smoke"), run_id)
    document_ids = [upload_and_ingest(client, path) for path in sample_paths]
    document_detail = unwrap(client.get(f"/api/v1/documents/{document_ids[0]}"))
    assert document_detail["chunk_count"] >= 1, document_detail

    search = unwrap(client.post("/api/v1/search", json={"query": "直营渠道", "top_k": 5}))
    assert len(search["items"]) >= 1, search

    qa = unwrap(client.post("/api/v1/qa/ask", json={"question": "收入增长原因是什么？", "top_k": 5}))
    assert qa["citations"], qa
    assert qa["answer_provider"] == "llm", qa
    assert qa["task"]["status"] == "completed", qa

    task_detail = unwrap(client.get(f"/api/v1/tasks/{qa['task']['id']}"))
    assert task_detail["status"] == "completed", task_detail
    assert task_detail["output_payload"]["citations_count"] == len(qa["citations"]), task_detail

    task_runs = unwrap(client.get(f"/api/v1/tasks/{qa['task']['id']}/runs"))
    assert len(task_runs["items"]) >= 1, task_runs
    assert task_runs["items"][0]["status"] == "completed", task_runs
    assert_citations_saved(qa["task"]["id"], len(qa["citations"]))

    task_execution = unwrap(
        client.post(
            "/api/v1/tasks",
            json={
                "task_type": "qa",
                "task_title": "Smoke QA",
                "input_text": "贵州茅台收入增长原因是什么？",
                "input_payload": {"question": "贵州茅台收入增长原因是什么？", "top_k": 3},
            },
        )
    )
    assert task_execution["task"]["status"] == "completed", task_execution
    assert task_execution["decision"]["agent"] == "research_agent", task_execution
    assert task_execution["result"]["task"]["status"] == "completed", task_execution
    assert_citations_saved(task_execution["task"]["id"], len(task_execution["result"]["citations"]))

    task_runs_from_tasks = unwrap(client.get(f"/api/v1/tasks/{task_execution['task']['id']}/runs"))
    run_names = {item["run_name"] for item in task_runs_from_tasks["items"]}
    assert "orchestrator_agent.route" in run_names, task_runs_from_tasks
    assert "research_agent" in run_names, task_runs_from_tasks
    task_list = unwrap(client.get("/api/v1/tasks", params={"task_type": "qa", "status": "completed"}))
    assert any(item["id"] == task_execution["task"]["id"] for item in task_list["items"]), task_list

    pending_task = unwrap(
        client.post(
            "/api/v1/tasks",
            json={"task_type": "qa", "task_title": "Pending QA", "input_text": "收入增长原因", "execute": False},
        )
    )
    assert pending_task["status"] == "pending", pending_task
    executed_task = unwrap(
        client.post(
            "/api/v1/qa/ask",
            json={"question": "贵州茅台收入增长原因是什么？", "top_k": 3, "task_id": pending_task["id"]},
        )
    )
    assert executed_task["task"]["id"] == pending_task["id"], executed_task
    assert executed_task["task"]["status"] == "completed", executed_task
    assert_citations_saved(pending_task["id"], len(executed_task["citations"]))

    routed = unwrap(
        client.post(
            "/api/v1/agents/route",
            json={"task_type": "qa", "question": "贵州茅台收入增长原因是什么？", "top_k": 3},
        )
    )
    assert routed["decision"]["agent"] == "research_agent", routed
    assert routed["result"]["task"]["status"] == "completed", routed

    memo = unwrap(client.post("/api/v1/qa/memo", json={"topic": "贵州茅台收入增长原因", "top_k": 3, "company_code": "600519"}))
    assert memo["asset"]["asset_type"] == "memo", memo
    assert "##" in memo["content_markdown"], memo
    assert_asset_citations_saved(memo["asset"]["id"], len(memo["citations"]))
    assert memo["asset"]["sources"], memo

    memo_sources = unwrap(client.get(f"/api/v1/assets/{memo['asset']['id']}/sources"))
    assert len(memo_sources["items"]) >= 1, memo_sources

    tag = unwrap(
        client.post(
            "/api/v1/tags",
            json={
                "tag_code": f"smoke-tag-{run_id}",
                "tag_name": f"Smoke Tag {run_id}",
                "tag_type": "topic",
            },
        )
    )

    daily = unwrap(client.post("/api/v1/qa/daily-report", json={"topic": "白酒行业今日重点", "top_k": 3}))
    assert daily["asset"]["asset_type"] == "daily_report", daily
    assert "## 市场概览" in daily["content_markdown"], daily
    assert "## 公告跟踪" in daily["content_markdown"], daily
    assert "## 新闻动态" in daily["content_markdown"], daily
    assert "## 研究观点" in daily["content_markdown"], daily
    assert_asset_citations_saved(daily["asset"]["id"], len(daily["citations"]))

    routed_daily = unwrap(
        client.post(
            "/api/v1/agents/route",
            json={"task_type": "daily_report", "topic": "白酒行业今日重点", "top_k": 3, "export_format": "docx"},
        )
    )
    assert routed_daily["decision"]["agent"] == "report_agent", routed_daily
    assert routed_daily["result"]["report_type"] == "daily_report", routed_daily
    assert routed_daily["result"]["export"]["format"] == "docx", routed_daily
    assert routed_daily["result"]["export"]["file_name"].endswith(".docx"), routed_daily
    assert Path(routed_daily["result"]["export"]["file_path"]).exists(), routed_daily

    asset = unwrap(
        client.post(
            "/api/v1/assets",
            json={
                "asset_type": "memo",
                "title": "Smoke Memo",
                "content_markdown": "# Smoke Memo\n\n收入增长来自直营渠道。",
                "company_code": "600519",
                "sources": [{"source_type": "document", "source_ref_id": document_ids[0], "source_id_text": "smoke-doc"}],
                "tag_ids": [tag["id"]],
            },
        )
    )
    assert asset["sources"], asset
    assert len(asset["tags"]) == 1, asset
    asset_list = unwrap(client.get("/api/v1/assets", params={"company_code": "600519", "tag_id": tag["id"], "keyword": "Smoke"}))
    assert any(item["id"] == asset["id"] for item in asset_list["items"]), asset_list

    updated_asset = unwrap(
        client.put(
            f"/api/v1/assets/{asset['id']}",
            json={
                "content_markdown": "# Smoke Memo\n\n收入增长来自直营渠道，且产品结构持续优化。",
                "summary": "更新后的人工修订版本",
                "change_note": "manual revision",
                "sources": [{"source_type": "document", "source_ref_id": document_ids[1], "source_id_text": "smoke-doc-2"}],
                "tag_ids": [tag["id"]],
            },
        )
    )
    assert updated_asset["version"] == 2, updated_asset
    assert "产品结构持续优化" in updated_asset["content_markdown"], updated_asset

    revisions = unwrap(client.get(f"/api/v1/assets/{asset['id']}/revisions"))
    assert len(revisions["items"]) >= 2, revisions
    exported = unwrap(client.post(f"/api/v1/assets/{asset['id']}/export", json={"format": "markdown"}))
    assert exported["file_name"].endswith(".md"), exported
    assert exported["file_name"].startswith("memo_600519_"), exported
    assert exported["download_path"] == exported["file_path"], exported
    assert Path(exported["file_path"]).exists(), exported

    print(
        {
            "status": "ok",
            "documents": document_ids,
            "search_hits": len(search["items"]),
            "citations": len(qa["citations"]),
            "qa_task_id": qa["task"]["id"],
            "task_api_task_id": task_execution["task"]["id"],
            "manual_task_id": executed_task["task"]["id"],
            "memo_asset_id": memo["asset"]["id"],
            "daily_asset_id": daily["asset"]["id"],
            "asset_id": asset["id"],
        }
    )


if __name__ == "__main__":
    main()
